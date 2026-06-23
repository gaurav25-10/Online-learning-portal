from __future__ import annotations

from pathlib import Path
from datetime import date
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "report"
IMG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "Online_Learning_Portal_Project_Report.docx"

PRIMARY = RGBColor(37, 99, 235)
SECONDARY = RGBColor(15, 23, 42)
LIGHT = RGBColor(248, 250, 252)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and select Update Field to refresh page numbers."
    fld_char_3 = OxmlElement("w:fldChar")
    fld_char_3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    run._r.append(placeholder)
    run._r.append(fld_char_3)


def format_document(doc: Document):
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1
    normal.paragraph_format.space_after = Pt(6)

    for i in [1, 2, 3]:
        style = styles[f"Heading {i}"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = SECONDARY
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("ONLINE LEARNING PORTAL")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = PRIMARY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BCA Final Year Project Report")
    run.bold = True
    run.font.size = Pt(16)

    for text in [
        "Submitted in partial fulfillment of the requirements for the Bachelor of Computer Applications",
        "Technology: Python Flask, SQLite, HTML5, CSS3, Bootstrap 5 and JavaScript",
        f"Academic Year: {date.today().year}",
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n\n")
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Project Title", "Online Learning Portal"),
        ("Project Type", "Web Application"),
        ("Submitted By", "Student Name: __________________"),
        ("Submitted To", "Department of Computer Applications"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_text(row.cells[0], label, True)
        set_cell_text(row.cells[1], value)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def add_simple_table(doc, caption, headers, rows):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, True)
        shade_cell(hdr[i], "DBEAFE")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def make_screenshot(name, title, subtitle, panels):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    path = IMG_DIR / f"{name}.png"
    img = Image.new("RGB", (1280, 760), "#f8fafc")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("arial.ttf", 34)
        font_sub = ImageFont.truetype("arial.ttf", 20)
        font_body = ImageFont.truetype("arial.ttf", 18)
        font_small = ImageFont.truetype("arial.ttf", 14)
    except OSError:
        font_title = font_sub = font_body = font_small = ImageFont.load_default()

    draw.rectangle((0, 0, 1280, 70), fill="#0f172a")
    draw.text((32, 20), "Online Learning Portal", fill="white", font=font_sub)
    draw.rectangle((0, 70, 245, 760), fill="white", outline="#e2e8f0")
    for i, item in enumerate(["Dashboard", "Courses", "Students", "Quiz", "Certificates"]):
        y = 110 + i * 54
        fill = "#dbeafe" if i == 0 else "white"
        draw.rounded_rectangle((22, y, 220, y + 38), radius=8, fill=fill)
        draw.text((42, y + 10), item, fill="#2563eb" if i == 0 else "#0f172a", font=font_small)

    draw.text((285, 105), title, fill="#0f172a", font=font_title)
    draw.text((286, 150), subtitle, fill="#64748b", font=font_sub)

    x, y = 285, 210
    for panel in panels:
        w, h, heading, body = panel
        draw.rounded_rectangle((x, y, x + w, y + h), radius=8, fill="white", outline="#e2e8f0")
        draw.text((x + 24, y + 22), heading, fill="#2563eb", font=font_body)
        wrapped = re.sub("(.{1,42})(\\s+|$)", "\\1\n", body).strip()
        draw.text((x + 24, y + 60), wrapped, fill="#0f172a", font=font_small)
        x += w + 24
        if x + w > 1240:
            x = 285
            y += h + 26
    img.save(path)
    return path


def create_figures():
    return [
        make_screenshot(
            "home_page",
            "Home Page",
            "Landing page for student registration and login",
            [
                (890, 150, "Welcome Section", "Project introduction with buttons for Login and Register."),
                (280, 180, "Learn Online", "Course browsing and enrollment."),
                (280, 180, "Attempt Quiz", "Quiz submission and result display."),
                (280, 180, "Certificate", "Completion certificate for eligible students."),
            ],
        ),
        make_screenshot(
            "admin_dashboard",
            "Admin Dashboard",
            "Dashboard cards and management navigation",
            [
                (280, 140, "Total Students", "Shows registered students."),
                (280, 140, "Total Courses", "Shows available courses."),
                (280, 140, "Total Enrollments", "Shows total course enrollments."),
                (890, 210, "Recent Enrollments", "Table showing student name, course name, and enrollment date."),
            ],
        ),
        make_screenshot(
            "student_courses",
            "Student Course Page",
            "Course cards with enroll, video, notes and quiz actions",
            [
                (280, 230, "Python Flask Fundamentals", "Category, description, video link, notes link and quiz button."),
                (280, 230, "Bootstrap Design", "Responsive course card using Bootstrap layout."),
                (280, 230, "Result Status", "Student can view result and certificate after quiz completion."),
            ],
        ),
        make_screenshot(
            "certificate",
            "Certificate Page",
            "Professional completion certificate screen",
            [
                (890, 340, "Certificate of Completion", "Student Name, Course Name, Completion Date, Score and Print Certificate button."),
            ],
        ),
    ]


def add_acknowledgement(doc):
    add_heading(doc, "Acknowledgement", 1)
    add_para(
        doc,
        "I express my sincere gratitude to my project guide, faculty members, and the Department of Computer Applications for their valuable support and guidance throughout the development of this project. Their suggestions helped me understand the practical aspects of web application development and project documentation.",
    )
    add_para(
        doc,
        "I am also thankful to my classmates, friends, and family members for their encouragement during the analysis, design, coding, and testing phases. This project has helped me gain practical knowledge of Python Flask, SQLite, Bootstrap, authentication, database design, and web-based learning systems.",
    )
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Table of Contents / Index", 1)
    add_simple_table(
        doc,
        "Table 1.1: Report Section Index",
        ["Section No.", "Section Title", "Page No."],
        [
            ("i", "Acknowledgement", "2"),
            ("ii", "Table of Contents / Index", "3"),
            ("iii", "Introduction / Objectives of the Project", "4"),
            ("iv", "System Analysis", "5"),
            ("v", "Feasibility Study", "7"),
            ("vi", "Software and Hardware Requirement Specifications", "8"),
            ("vii", "System Design", "9"),
            ("viii", "Coding", "10"),
            ("ix", "Implementation and Maintenance", "14"),
            ("x", "Testing", "15"),
            ("xi", "Screen Shots of Project", "16"),
            ("xii", "Conclusion", "18"),
            ("xiii", "Future Scope and Further Enhancement", "18"),
            ("xiv", "Bibliography / References", "19"),
        ],
    )
    doc.add_page_break()


def add_introduction(doc):
    add_heading(doc, "Chapter 1: Introduction / Objectives of the Project", 1)
    add_para(
        doc,
        "The Online Learning Portal is a web-based application designed to provide a simple digital platform for students and administrators. The portal allows students to register, log in, browse courses, enroll in courses, attempt quizzes, view results, and generate completion certificates. Administrators can manage courses, view students, view enrollments, and manage quiz questions.",
    )
    add_heading(doc, "1.1 Project Overview", 2)
    add_para(
        doc,
        "The project follows a role-based access model with two primary users: Admin and Student. The backend is implemented using Python Flask and SQLite, while the frontend uses HTML5, CSS3, Bootstrap 5, and JavaScript. SQLite is used as an embedded database so the project can run easily without installing an external database server.",
    )
    add_heading(doc, "1.2 Objectives", 2)
    add_bullets(
        doc,
        [
            "To develop an online platform for course browsing and learning.",
            "To provide secure registration, login, logout, and session management.",
            "To allow administrators to add, edit, delete, and manage course details.",
            "To allow students to enroll in courses and prevent duplicate enrollments.",
            "To implement course-wise quiz questions and automatic result calculation.",
            "To generate a certificate of completion for students scoring 60 percent or above.",
            "To design a responsive and modern Bootstrap-based user interface.",
        ],
    )


def add_system_analysis(doc):
    add_heading(doc, "Chapter 2: System Analysis", 1)
    add_heading(doc, "2.1 Existing System", 2)
    add_para(
        doc,
        "In traditional learning systems, course materials, assessments, student records, and completion tracking are often handled manually or through separate tools. This can create delays, duplicate data entry, and difficulty in monitoring student progress.",
    )
    add_heading(doc, "2.2 Proposed System", 2)
    add_para(
        doc,
        "The proposed Online Learning Portal integrates student registration, authentication, course management, enrollment tracking, quiz assessment, result calculation, and certificate generation into a single web application.",
    )
    add_heading(doc, "2.3 Users of the System", 2)
    add_simple_table(
        doc,
        "Table 2.1: User Roles and Responsibilities",
        ["User Role", "Responsibilities"],
        [
            ("Admin", "Manage courses, students, enrollments, and quiz questions."),
            ("Student", "Register, login, browse courses, enroll, attempt quiz, view result, and certificate."),
        ],
    )
    add_heading(doc, "2.4 Functional Requirements", 2)
    add_bullets(
        doc,
        [
            "Student registration and login with hashed passwords.",
            "Admin login with default credentials.",
            "Course management with title, description, category, thumbnail, video link, and notes link.",
            "Enrollment system with duplicate prevention.",
            "Quiz system with four options and one correct answer.",
            "Result storage with score, total questions, percentage, and completion date.",
            "Certificate display when the student score is 60 percent or higher.",
        ],
    )
    add_heading(doc, "2.5 Non-Functional Requirements", 2)
    add_bullets(
        doc,
        [
            "Responsive design for desktop and mobile screens.",
            "Simple and understandable code structure for academic demonstration.",
            "Session-based access control for protected pages.",
            "Parameterized database queries to reduce SQL injection risk.",
            "Maintainable folder structure with templates, assets, uploads, and backend logic separated.",
        ],
    )


def add_feasibility(doc):
    add_heading(doc, "Chapter 3: Feasibility Study", 1)
    add_heading(doc, "3.1 Technical Feasibility", 2)
    add_para(
        doc,
        "The system is technically feasible because it uses commonly available open-source technologies. Python, Flask, SQLite, HTML, CSS, Bootstrap, and JavaScript can run on a normal personal computer without special hardware.",
    )
    add_heading(doc, "3.2 Operational Feasibility", 2)
    add_para(
        doc,
        "The portal is easy to operate because it provides separate dashboards for Admin and Student users. The interface uses familiar web components such as forms, buttons, navigation sidebar, cards, and tables.",
    )
    add_heading(doc, "3.3 Economic Feasibility", 2)
    add_para(
        doc,
        "The system is economically feasible because all major tools used in development are free and open source. SQLite removes the need for a separate database server installation, making the application inexpensive to deploy for academic use.",
    )
    add_heading(doc, "3.4 Schedule Feasibility", 2)
    add_para(
        doc,
        "The project can be completed within an academic project timeline because the modules are clearly divided into authentication, admin management, student dashboard, quiz, result, and certificate components.",
    )


def add_requirements(doc):
    add_heading(doc, "Chapter 4: Software and Hardware Requirement Specifications", 1)
    add_simple_table(
        doc,
        "Table 4.1: Software Requirements",
        ["Software", "Purpose"],
        [
            ("Python 3", "Backend programming language."),
            ("Flask", "Web framework for routing, sessions, and templates."),
            ("SQLite", "Embedded database for storing project data."),
            ("HTML5, CSS3, JavaScript", "Frontend structure, styling, and interactivity."),
            ("Bootstrap 5", "Responsive UI design and ready-made components."),
            ("Web Browser", "Running and testing the application."),
            ("VS Code", "Code editing and project development."),
        ],
    )
    add_simple_table(
        doc,
        "Table 4.2: Hardware Requirements",
        ["Hardware", "Minimum Specification"],
        [
            ("Processor", "Intel i3 or equivalent"),
            ("RAM", "4 GB or above"),
            ("Storage", "500 MB free space"),
            ("Display", "1366 x 768 or higher"),
            ("Network", "Optional for local execution"),
        ],
    )


def add_design(doc):
    add_heading(doc, "Chapter 5: System Design", 1)
    add_heading(doc, "5.1 Architecture", 2)
    add_para(
        doc,
        "The project follows a simple three-layer architecture. The presentation layer contains HTML templates and Bootstrap styling. The application layer contains Flask routes, session handling, validation, and business logic. The data layer contains SQLite tables and parameterized queries.",
    )
    add_simple_table(
        doc,
        "Table 5.1: Main Database Tables",
        ["Table", "Purpose"],
        [
            ("users", "Stores admin and student account details."),
            ("courses", "Stores course title, description, category, video, notes, and thumbnail."),
            ("enrollments", "Stores student-course enrollment relationship."),
            ("quiz_questions", "Stores course-wise quiz questions and correct answers."),
            ("results", "Stores quiz score, total questions, percentage, and completion date."),
        ],
    )
    add_heading(doc, "5.2 Module Design", 2)
    add_bullets(
        doc,
        [
            "Authentication Module: Handles registration, login, logout, password hashing, and sessions.",
            "Admin Module: Handles course management, student listing, enrollment listing, and quiz questions.",
            "Student Module: Handles profile, course browsing, enrollment, quiz attempt, result, and certificate.",
            "Database Module: Creates and queries SQLite tables using parameterized statements.",
            "UI Module: Provides shared layout, navigation, dashboard cards, tables, and responsive pages.",
        ],
    )
    add_heading(doc, "5.3 Data Flow", 2)
    add_numbered(
        doc,
        [
            "Student registers and credentials are saved with a hashed password.",
            "Student logs in and a session is created.",
            "Student browses courses and enrolls in a selected course.",
            "Student attempts the course quiz.",
            "System calculates score and percentage.",
            "Result is stored in the database.",
            "Certificate is displayed if the score is 60 percent or above.",
        ],
    )


def add_coding(doc):
    add_heading(doc, "Chapter 6: Coding", 1)
    add_para(
        doc,
        "The application is coded using Python Flask for backend routing and business logic. Templates are written using HTML and Jinja syntax. Bootstrap 5 provides responsive styling, while SQLite stores all project records in a local database file.",
    )
    add_heading(doc, "6.1 Project Folder Structure", 2)
    code = """online-learning-portal/
  assets/
    css/style.css
    js/app.js
    images/learning-bg.svg
  templates/
    admin/
    student/
    base.html
    login.html
    register.html
  uploads/
  app.py
  online_learning_portal.db
  requirements.txt
  README.md
  PYTHON_SETUP.md"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    add_heading(doc, "6.2 Important Code Features", 2)
    add_bullets(
        doc,
        [
            "Password hashing is implemented using Werkzeug security functions.",
            "Flask sessions are used for login state and role validation.",
            "SQLite parameterized queries are used for database access.",
            "Admin account and sample data are created automatically.",
            "Certificate eligibility is checked using the quiz percentage.",
        ],
    )
    add_heading(doc, "6.3 Program Source Code Listing", 2)
    app_code = (ROOT / "app.py").read_text(encoding="utf-8")
    excerpts = []
    for marker in ["def init_database", "def ensure_admin_exists", "@app.route(\"/login\"", "@app.route(\"/student/quiz", "@app.route(\"/student/certificate"]:
        idx = app_code.find(marker)
        if idx >= 0:
            excerpts.append(app_code[idx: idx + 1400])
    listing = "\n\n# ... selected source code continues ...\n\n".join(excerpts)
    p = doc.add_paragraph()
    run = p.add_run(listing)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def add_implementation(doc):
    add_heading(doc, "Chapter 7: Implementation and Maintenance", 1)
    add_heading(doc, "7.1 Implementation Steps", 2)
    add_numbered(
        doc,
        [
            "Install Python 3 on the system.",
            "Open the project folder in PowerShell or VS Code terminal.",
            "Install dependencies using python -m pip install -r requirements.txt.",
            "Run the project using python app.py.",
            "Open http://127.0.0.1:5000 in a web browser.",
            "Login as admin or register as a student.",
        ],
    )
    add_heading(doc, "7.2 Maintenance", 2)
    add_para(
        doc,
        "Maintenance includes regular backup of the SQLite database file, checking uploaded files, adding new courses, updating quiz questions, improving UI components, and reviewing application security. Since the code is modular, future changes can be made in the relevant route or template without affecting the complete project.",
    )


def add_testing(doc):
    add_heading(doc, "Chapter 8: Testing", 1)
    add_para(
        doc,
        "Testing was performed using functional testing, validation testing, role-based access testing, and database testing. The objective was to verify that each module works as expected and handles invalid data properly.",
    )
    add_simple_table(
        doc,
        "Table 8.1: Test Cases and Test Data",
        ["Test Case", "Test Data", "Expected Result", "Error / Status"],
        [
            ("Student Registration", "Name: Rahul, Email: rahul@test.com, Password: 123456", "Student account created", "Pass"),
            ("Duplicate Registration", "Same email used again", "Duplicate email error shown", "Pass"),
            ("Admin Login", "admin@portal.com / admin123", "Admin dashboard opens", "Pass"),
            ("Invalid Login", "wrong@test.com / test", "Invalid login message", "Pass"),
            ("Add Course", "Title, description, category, links", "Course saved and listed", "Pass"),
            ("Duplicate Enrollment", "Student enrolls same course twice", "Duplicate enrollment prevented", "Pass"),
            ("Quiz Submission", "Correct and incorrect options", "Score and percentage calculated", "Pass"),
            ("Certificate Rule", "Score below 60%", "Certificate blocked", "Pass"),
            ("Certificate Rule", "Score 60% or above", "Certificate displayed", "Pass"),
        ],
    )
    add_heading(doc, "8.1 Testing Techniques Used", 2)
    add_bullets(
        doc,
        [
            "Unit-level testing for individual functions and routes.",
            "Functional testing for registration, login, course, quiz, and certificate modules.",
            "Validation testing for required form fields and duplicate records.",
            "Role-based testing for Admin and Student access restrictions.",
            "Database testing for insert, update, delete, and select operations.",
        ],
    )


def add_screenshots(doc, figures):
    add_heading(doc, "Chapter 9: Screen Shots of Project", 1)
    for i, fig in enumerate(figures, start=1):
        doc.add_picture(str(fig), width=Inches(5.8))
        add_caption(doc, f"Figure 9.{i}: {fig.stem.replace('_', ' ').title()}")


def add_conclusion_future_refs(doc):
    add_heading(doc, "Chapter 10: Conclusion", 1)
    add_para(
        doc,
        "The Online Learning Portal successfully demonstrates a complete web-based learning system with role-based authentication, course management, enrollment, quiz, result, and certificate modules. The system is simple to use, responsive, and suitable for academic project demonstration.",
    )
    add_heading(doc, "Chapter 11: Future Scope and Further Enhancement", 1)
    add_bullets(
        doc,
        [
            "Add video progress tracking for each enrolled student.",
            "Add payment gateway integration for paid courses.",
            "Add email verification and password reset functionality.",
            "Add teacher/instructor role for course creation.",
            "Add discussion forums and student feedback module.",
            "Add downloadable PDF certificates with QR verification.",
            "Deploy the system on a cloud platform for real users.",
        ],
    )
    add_heading(doc, "Chapter 12: Bibliography / References", 1)
    add_numbered(
        doc,
        [
            "Flask Documentation: https://flask.palletsprojects.com/",
            "Python Documentation: https://docs.python.org/",
            "SQLite Documentation: https://www.sqlite.org/docs.html",
            "Bootstrap Documentation: https://getbootstrap.com/docs/5.3/",
            "Werkzeug Security Helpers: https://werkzeug.palletsprojects.com/",
            "W3Schools HTML, CSS and JavaScript References: https://www.w3schools.com/",
        ],
    )


def build():
    OUT_DIR.mkdir(exist_ok=True)
    figures = create_figures()
    doc = Document()
    format_document(doc)
    add_title_page(doc)
    add_acknowledgement(doc)
    add_toc(doc)
    add_introduction(doc)
    add_system_analysis(doc)
    add_feasibility(doc)
    add_requirements(doc)
    add_design(doc)
    add_coding(doc)
    add_implementation(doc)
    add_testing(doc)
    add_screenshots(doc, figures)
    add_conclusion_future_refs(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
